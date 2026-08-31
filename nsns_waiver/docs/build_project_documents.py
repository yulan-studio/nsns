from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT_DIR = Path(__file__).resolve().parent
SPEC_PATH = OUTPUT_DIR / "NSNS_Waiver_Functional_Specification.docx"
DESIGN_PATH = OUTPUT_DIR / "NSNS_Waiver_Design_and_Implementation.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
LIGHT_BLUE = "E8EEF5"
PALE_BLUE = "F4F7FB"
GRAY = "666666"
LIGHT_GRAY = "D9E2F3"
WHITE = "FFFFFF"
BLACK = "000000"
AMBER = "FFF3CD"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_inches: list[float]) -> None:
    widths_dxa = [round(width * 1440) for width in widths_inches]
    total = sum(widths_dxa)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    old_grid = table._tbl.tblGrid
    for child in list(old_grid):
        old_grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        old_grid.append(grid_col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_repeat_and_keep(paragraph, keep_next=False) -> None:
    paragraph.paragraph_format.keep_together = True
    paragraph.paragraph_format.keep_with_next = keep_next


def set_font(run, name="Calibri", size=None, color=None, bold=None, italic=None) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_page_number(paragraph) -> None:
    paragraph.add_run("Page ")
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    value = OxmlElement("w:t")
    value.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for element in (begin, instruction, separate, value, end):
        run._r.append(element)


def set_paragraph_bottom_border(paragraph, color=BLUE, size=12) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def configure_document(title: str, subtitle: str, document_type: str) -> Document:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_tokens = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, DARK_BLUE, 10, 5),
    }
    for style_name, (size, color, before, after) in heading_tokens.items():
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    header = section.header
    header_p = header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    left = header_p.add_run(document_type.upper())
    set_font(left, size=8.5, color=GRAY, bold=True)
    right = header_p.add_run("  |  NSNS Waiver Application")
    set_font(right, size=8.5, color=GRAY)

    footer = section.footer
    footer_p = footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_page_number(footer_p)
    for run in footer_p.runs:
        set_font(run, size=8.5, color=GRAY)

    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_before = Pt(8)
    kicker.paragraph_format.space_after = Pt(2)
    run = kicker.add_run(document_type.upper())
    set_font(run, size=9, color=BLUE, bold=True)

    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(4)
    title_run = title_p.add_run(title)
    set_font(title_run, size=25, color=DARK_BLUE, bold=True)

    subtitle_p = doc.add_paragraph()
    subtitle_p.paragraph_format.space_after = Pt(12)
    subtitle_run = subtitle_p.add_run(subtitle)
    set_font(subtitle_run, size=12.5, color=GRAY)
    set_paragraph_bottom_border(subtitle_p)

    add_metadata_table(
        doc,
        [
            ("System", "NSNS Waiver Application"),
            ("Status", "As implemented"),
            ("Technology", "ASP.NET Core Razor Pages / .NET 10 / MySQL / Dapper"),
            ("Prepared", date.today().isoformat()),
        ],
    )
    doc.core_properties.title = title
    doc.core_properties.subject = subtitle
    doc.core_properties.author = "NSNS"
    doc.core_properties.keywords = "NSNS, waiver, specification, design, implementation"
    return doc


def add_metadata_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value
        set_cell_shading(cells[0], LIGHT_BLUE)
        for run in cells[0].paragraphs[0].runs:
            set_font(run, size=9.5, color=DARK_BLUE, bold=True)
        for run in cells[1].paragraphs[0].runs:
            set_font(run, size=9.5)
    set_table_geometry(table, [1.181, 5.319])
    doc.add_paragraph()


def add_heading(doc: Document, text: str, level=1) -> None:
    p = doc.add_paragraph(text, style=f"Heading {level}")
    set_repeat_and_keep(p, keep_next=True)


def add_body(doc: Document, text: str, bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        lead = p.add_run(bold_prefix)
        set_font(lead, bold=True, color=DARK_BLUE)
        p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        set_repeat_and_keep(p)


def add_numbers(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(item, style="List Number")
        set_repeat_and_keep(p)


def add_callout(doc: Document, label: str, text: str, fill=PALE_BLUE) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    label_run = p.add_run(f"{label}: ")
    set_font(label_run, bold=True, color=DARK_BLUE)
    p.add_run(text)
    set_table_geometry(table, [6.5])


def add_table(
    doc: Document,
    headers: list[str],
    rows: list[list[str]],
    widths: list[float],
    font_size=9,
) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = header
        set_cell_shading(cell, LIGHT_BLUE)
        for run in cell.paragraphs[0].runs:
            set_font(run, size=font_size, color=DARK_BLUE, bold=True)
    repeat_table_header(table.rows[0])
    for row_data in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row_data):
            cells[index].text = value
            for p in cells[index].paragraphs:
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.line_spacing = 1.05
                for run in p.runs:
                    set_font(run, size=font_size)
    set_table_geometry(table, widths)


def add_contents(doc: Document, entries: list[str]) -> None:
    add_heading(doc, "Document contents", 1)
    for entry in entries:
        doc.add_paragraph(entry, style="List Bullet")
    doc.add_page_break()


def add_requirement_table(doc: Document, requirements: list[tuple[str, str, str]]) -> None:
    add_table(
        doc,
        ["ID", "Requirement", "Verification"],
        [[identifier, requirement, verification] for identifier, requirement, verification in requirements],
        [0.65, 4.85, 1.0],
        font_size=8.5,
    )


def build_specification() -> None:
    doc = configure_document(
        "Functional Specification",
        "Business and system requirements for event waiver collection",
        "System Specification",
    )
    add_callout(
        doc,
        "Purpose",
        "Defines what the NSNS Waiver Application must do, the rules it enforces, "
        "the information it stores, and the acceptance conditions for operation.",
    )
    add_contents(
        doc,
        [
            "1. Purpose and scope",
            "2. Stakeholders and operating context",
            "3. Functional requirements",
            "4. Business and validation rules",
            "5. Data requirements",
            "6. Non-functional requirements",
            "7. Acceptance criteria",
            "8. Assumptions, exclusions, and dependencies",
        ],
    )

    add_heading(doc, "1. Purpose and scope")
    add_body(
        doc,
        "The application lets a customer open an event-specific link, enter contact "
        "information, add family members, review and sign a waiver, choose media-release "
        "consent, and receive confirmation. It also queues a business-owner notification "
        "and provides a protected administrator list of recent submissions.",
    )
    add_heading(doc, "1.1 In scope", 2)
    add_bullets(
        doc,
        [
            "Public waiver form selected by the event query-string parameter.",
            "Server-side event validation against configuration.",
            "Contact details, electronic signature, family members, and optional media consent.",
            "Transactional MySQL persistence and email outbox creation.",
            "Customer confirmation and business-owner notification emails.",
            "Cookie-protected administrator login and paginated submission list.",
            "Docker-compatible deployment with Railway-oriented configuration.",
        ],
    )
    add_heading(doc, "1.2 Out of scope", 2)
    add_bullets(
        doc,
        [
            "Self-service event management or an Events database table.",
            "Waiver versioning, agreement snapshots, agreement hashes, or metadata tables.",
            "Duplicate-submission prevention.",
            "Search, export, customer accounts, payment collection, or email-template administration.",
            "Automated legal approval of the agreement content.",
        ],
    )

    add_heading(doc, "2. Stakeholders and operating context")
    add_table(
        doc,
        ["Actor", "Goal", "Authorized capabilities"],
        [
            ["Customer / signer", "Submit a waiver for an event.", "View event and agreement; enter data; add family members; sign; choose media consent; submit."],
            ["Business owner", "Receive notice and review submissions.", "Receive notification email; sign in; view and sort recent submissions; sign out."],
            ["Operator / maintainer", "Configure and run the service.", "Configure events, owner email, SMTP, admin credentials, database, and deployment."],
            ["Email worker", "Deliver queued messages reliably.", "Poll eligible outbox rows; send; mark sent; retry; abandon after maximum attempts."],
        ],
        [1.25, 2.1, 3.15],
        font_size=8.7,
    )
    add_callout(
        doc,
        "Primary entry point",
        "Customers use a URL such as https://waiver.example.com/?event=summer-camp-2026. "
        "The event code is never accepted solely from the browser; it is resolved against server configuration.",
    )

    add_heading(doc, "3. Functional requirements")
    add_heading(doc, "3.1 Public event and form", 2)
    add_requirement_table(
        doc,
        [
            ("FR-001", "The system shall read the event code from the event query-string parameter.", "Page test"),
            ("FR-002", "The system shall resolve the event code against configured allowed events using trimmed, case-insensitive matching.", "Service test"),
            ("FR-003", "The system shall show an event-not-found state when the event is missing or invalid.", "Page test"),
            ("FR-004", "The form shall collect first name, last name, optional WeChat name, email, and phone number.", "UI review"),
            ("FR-005", "The form shall allow zero to ten family members, each with first name, last name, and optional relationship.", "Page and service tests"),
            ("FR-006", "The system shall render the agreement from Content/waiver-agreement.html and disable submission when it is unavailable or contains the placeholder marker.", "Provider/page tests"),
            ("FR-007", "The customer shall be required to accept the liability waiver before submission.", "Page/service tests"),
            ("FR-008", "The customer shall enter a legal name as an electronic signature.", "Validation test"),
            ("FR-009", "The media-release checkbox shall be optional, selected by default, and may be cleared before submission.", "Page/service tests"),
        ],
    )

    add_heading(doc, "3.2 Submission processing", 2)
    add_requirement_table(
        doc,
        [
            ("FR-010", "The server shall revalidate all posted values regardless of browser validation.", "Service tests"),
            ("FR-011", "The server shall generate the signed timestamp in UTC and a UUID submission reference.", "Service test"),
            ("FR-012", "The system shall store the configured event code and event name with each submission.", "Service/repository tests"),
            ("FR-013", "The system shall allow duplicate submissions.", "Service test"),
            ("FR-014", "The submission, family members, customer email, and owner email shall be created in one database transaction.", "Repository test"),
            ("FR-015", "A successful submission shall redirect to a confirmation page showing the event, reference, and signed time.", "Page test"),
        ],
    )

    add_heading(doc, "3.3 Email delivery", 2)
    add_requirement_table(
        doc,
        [
            ("FR-016", "The system shall queue a professional confirmation email to the customer.", "Service test"),
            ("FR-017", "The system shall queue an owner notification containing customer, signature, family, event, and media-consent details.", "Service test"),
            ("FR-018", "The owner email shall link to https://waiver.nsns.ca/Admin/Submissions.", "Service test"),
            ("FR-019", "A background worker shall process eligible outbox messages in configured batches.", "Processor test"),
            ("FR-020", "Failed email delivery shall use exponential retry and become abandoned after the configured maximum attempts.", "Processor test"),
        ],
    )

    add_heading(doc, "3.4 Administrator area", 2)
    add_requirement_table(
        doc,
        [
            ("FR-021", "Anonymous users shall be redirected to the administrator login page.", "Authentication configuration"),
            ("FR-022", "Valid configured credentials shall create an eight-hour sliding authentication cookie.", "Credential/page tests"),
            ("FR-023", "Post-login redirects shall accept only local return URLs.", "Code review"),
            ("FR-024", "The admin list shall show event, customer details, family members, media consent, signature, and UTC signed time.", "UI/repository review"),
            ("FR-025", "The list shall sort only by an allow-listed set of columns.", "Repository/page tests"),
            ("FR-026", "The list shall show at most 200 recent submissions, 20 per page, and no more than ten pages.", "Page/repository tests"),
            ("FR-027", "The administrator shall be able to sign out and delete the authentication cookie.", "Page test"),
        ],
    )

    add_heading(doc, "4. Business and validation rules")
    add_table(
        doc,
        ["Rule", "Definition"],
        [
            ["BR-001 Event", "A submission is valid only when its normalized event code exists in Waiver:Events and has a nonblank name."],
            ["BR-002 Required text", "First name, last name, email, phone, and signature are trimmed and required."],
            ["BR-003 Lengths", "Names and relationship: 100 characters; email: 320; phone: 40; signature: 200; IP address: 45; user agent: 500."],
            ["BR-004 Email", "Email must parse as a valid address; a lower-case normalized value is stored for lookup."],
            ["BR-005 Phone", "The display value is stored as entered after trimming; a digits-only normalized value must contain at least one digit."],
            ["BR-006 Family", "At most ten family members are processed. Excess members are rejected."],
            ["BR-007 Agreement", "Liability-waiver acceptance is mandatory; media-release consent is independent and optional."],
            ["BR-008 Media default", "New forms begin with media consent selected. The database column is NOT NULL with DEFAULT TRUE."],
            ["BR-009 Duplicates", "The same customer may submit multiple waivers for the same or different events."],
            ["BR-010 Time", "Stored signing, creation, attempt, retry, and sent timestamps use UTC."],
        ],
        [1.45, 5.05],
        font_size=8.7,
    )

    add_heading(doc, "5. Data requirements")
    add_heading(doc, "5.1 Waiver submissions", 2)
    add_body(
        doc,
        "The waiver_submissions table is the parent record. It stores the UUID reference, "
        "event identity, customer contact information, normalized lookup values, signature, "
        "required-waiver acceptance, media-release choice, UTC signing time, request metadata, "
        "and creation time.",
    )
    add_heading(doc, "5.2 Family members", 2)
    add_body(
        doc,
        "Each waiver_family_members row belongs to one submission through a foreign key with "
        "ON DELETE CASCADE. Family members are deliberately stored separately rather than embedded.",
    )
    add_heading(doc, "5.3 Email outbox", 2)
    add_body(
        doc,
        "Each email_outbox row belongs to a submission and stores recipient, subject, HTML body, "
        "delivery status, attempt count, retry schedule, sent time, and a safe error summary. "
        "Deleting a submission cascades to its queued messages.",
    )
    add_heading(doc, "5.4 Agreement storage", 2)
    add_body(
        doc,
        "The agreement remains a deployed HTML file. It is not copied into MySQL and has no "
        "version, snapshot, hash, or metadata record.",
    )

    add_heading(doc, "6. Non-functional requirements")
    add_table(
        doc,
        ["Area", "Requirement"],
        [
            ["Security", "Use anti-forgery protection, parameterized SQL, server-side validation, HTTP-only admin cookies, local-only redirects, and fixed-time credential comparison."],
            ["Privacy", "Avoid logging PII, passwords, connection strings, or SMTP secrets. Store request metadata only as required for troubleshooting."],
            ["Reliability", "Persist submissions and email messages atomically. Retry email independently so SMTP failure cannot lose a waiver."],
            ["Maintainability", "Keep UI in Razor Pages, business rules in services, SQL in repositories, and settings in typed options."],
            ["Performance", "Limit admin reads to the 200 newest records, page at 20 rows, clamp email batch size to 1-100, and index common lookup/order fields."],
            ["Accessibility", "Use semantic headings, labels, validation messages, keyboard-focusable agreement content, and status text for dynamic family-member controls."],
            ["Portability", "Run on .NET 10 with MySQL in Docker-compatible and Railway-compatible environments."],
        ],
        [1.35, 5.15],
        font_size=8.8,
    )

    add_heading(doc, "7. Acceptance criteria")
    add_numbers(
        doc,
        [
            "A valid configured event opens the correct form and an unknown event cannot be submitted.",
            "A customer can submit with zero or up to ten family members and can submit duplicates.",
            "Clearing the optional media checkbox still permits submission and stores Declined/false.",
            "A successful request creates exactly one submission, the supplied family rows, and two outbox rows in one transaction.",
            "The customer reaches confirmation only after persistence succeeds.",
            "The email worker marks successful messages Sent and schedules or abandons failed messages according to configuration.",
            "Unauthenticated users cannot view admin submissions.",
            "Admin sorting cannot inject arbitrary SQL, and pagination never exposes more than 200 recent submissions.",
            "Release build and automated tests complete without code errors.",
        ],
    )

    add_heading(doc, "8. Assumptions, exclusions, and dependencies")
    add_bullets(
        doc,
        [
            "A MySQL database exists and schema migrations are applied before the updated application starts.",
            "The operator provides valid Waiver, Email, Admin, and ConnectionStrings configuration.",
            "SMTP credentials permit relay from the configured sender address.",
            "The business owner is responsible for legal approval and encoding quality of the HTML agreement.",
            "Historical rows receive media_release_agreed = TRUE when migration 002 is applied because no earlier choice was recorded.",
            "MySQL integration tests require WAIVERAPP_TEST_MYSQL_CONNECTION; otherwise they are skipped.",
        ],
    )
    add_callout(
        doc,
        "Legal-content boundary",
        "This specification describes software behavior. It does not validate whether the waiver "
        "or media-release wording is legally sufficient.",
        fill=AMBER,
    )
    doc.save(SPEC_PATH)


def build_design() -> None:
    doc = configure_document(
        "Design and Implementation Guide",
        "Architecture, code organization, data flow, operations, and maintenance",
        "Technical Design",
    )
    add_callout(
        doc,
        "Audience",
        "Developers and operators maintaining the ASP.NET Core Razor Pages application, "
        "MySQL schema, email outbox, administrator area, and Railway deployment.",
    )
    add_contents(
        doc,
        [
            "1. Design goals and constraints",
            "2. Runtime architecture",
            "3. Project structure and components",
            "4. Core request workflows",
            "5. Data design and migrations",
            "6. Security and validation design",
            "7. Email outbox design",
            "8. Administrator design",
            "9. Configuration and deployment",
            "10. Testing and maintenance",
            "11. Function reference",
        ],
    )

    add_heading(doc, "1. Design goals and constraints")
    add_bullets(
        doc,
        [
            "Keep one ASP.NET Core Razor Pages project; do not introduce MVC controllers, Blazor, a SPA framework, or Entity Framework.",
            "Use dependency injection and async/await throughout request and persistence paths.",
            "Keep UI concerns in PageModels, business rules in services, and SQL in repositories.",
            "Use Dapper with MySqlConnector and parameterized SQL only.",
            "Store every timestamp in UTC and generate the signed time on the server.",
            "Preserve the agreement as a deployed HTML file with no database versioning or snapshots.",
            "Commit a waiver and its related family/email records atomically.",
        ],
    )

    add_heading(doc, "2. Runtime architecture")
    add_body(
        doc,
        "The application is a layered monolith. Razor Pages receive HTTP requests and map "
        "untrusted input into service requests. Services enforce business rules and construct "
        "domain records. Repositories execute parameterized Dapper SQL. MySQL stores submissions "
        "and the reliable email outbox. A hosted worker delivers queued messages through SMTP.",
    )
    add_table(
        doc,
        ["Layer", "Primary responsibility", "Key implementation"],
        [
            ["Presentation", "Public form, confirmation, error/privacy pages, admin login and list.", "Pages/*.cshtml and PageModel classes"],
            ["Business", "Event resolution, validation, normalization, record construction, email composition.", "WaiverSubmissionService"],
            ["Persistence", "Transactional writes and read projections.", "Repositories/*Repository.cs"],
            ["Infrastructure", "MySQL connections, SMTP, background polling, configuration.", "Data, SmtpEmailSender, EmailOutboxWorker"],
            ["Domain/data contracts", "Submission, family member, outbox, admin projection, requests/results.", "Models and Services records"],
        ],
        [1.15, 2.75, 2.6],
        font_size=8.7,
    )
    add_callout(
        doc,
        "Dependency direction",
        "Pages depend on service/repository abstractions. Services depend on repository abstractions. "
        "Repositories depend on the database connection factory. Infrastructure implementations are "
        "wired centrally in Program.cs.",
    )

    add_heading(doc, "3. Project structure and components")
    add_table(
        doc,
        ["Path", "Responsibility"],
        [
            ["Program.cs", "Registers configuration, authentication, repositories, services, SMTP, hosted worker, and middleware."],
            ["Pages/", "Razor markup and PageModels for public and administrator workflows."],
            ["Services/", "Business validation, agreement loading, credential checking, SMTP sending, and outbox processing."],
            ["Repositories/", "Dapper SQL and transactional persistence."],
            ["Models/", "Database-backed domain records and admin list projection."],
            ["Options/", "Strongly typed Waiver, Email/SMTP, and Admin configuration."],
            ["Data/", "MySQL connection factory and primary migration copies."],
            ["Database/Scripts/", "Operational copies of schema and upgrade SQL."],
            ["Content/", "Deployable waiver-agreement.html."],
            ["Tests/", "Unit, PageModel, SQL, processor, credential, and optional MySQL integration tests."],
            ["wwwroot/", "Application CSS/JavaScript and vendored Bootstrap/jQuery assets."],
        ],
        [1.45, 5.05],
        font_size=8.7,
    )

    add_heading(doc, "4. Core request workflows")
    add_heading(doc, "4.1 Display the waiver form", 2)
    add_numbers(
        doc,
        [
            "IndexModel.OnGetAsync receives the event query-string value.",
            "WaiverSubmissionService.FindEvent normalizes and resolves it against Waiver:Events.",
            "FileWaiverAgreementProvider.GetAsync reads Content/waiver-agreement.html.",
            "The PageModel renders the event, agreement, form fields, repeatable family controls, required waiver checkbox, and default-selected optional media checkbox.",
            "If the event is unknown, the page shows an event-not-found state. If the agreement is unavailable/unapproved, submission controls are disabled.",
        ],
    )
    add_heading(doc, "4.2 Submit a waiver", 2)
    add_numbers(
        doc,
        [
            "Razor model binding and DataAnnotations perform initial validation and anti-forgery validation occurs automatically.",
            "IndexModel.OnPostAsync re-resolves the event, reloads the agreement, verifies required agreement acceptance, and maps a SubmitWaiverRequest.",
            "WaiverSubmissionService.SubmitAsync trims, validates, normalizes, and accumulates all field errors.",
            "The service creates a server UTC timestamp, UUID reference, WaiverSubmission, family records, and two HTML-safe outbox messages.",
            "WaiverSubmissionRepository.CreateSubmissionAsync opens a transaction and inserts the parent, children, and outbox rows.",
            "After commit, the PageModel places safe confirmation values in TempData and redirects to /Confirmation.",
            "ConfirmationModel.OnGet consumes TempData and redirects stale/direct visits back to the index.",
        ],
    )
    add_heading(doc, "4.3 Review submissions", 2)
    add_numbers(
        doc,
        [
            "LoginModel validates configured credentials and creates the cookie principal.",
            "SubmissionsModel parses allow-listed sort values and clamps the page number.",
            "AdminSubmissionRepository selects from the 200 newest records and aggregates family members with GROUP_CONCAT.",
            "The Razor table displays 20 rows per page and uses links to toggle sort direction.",
            "Logout deletes the cookie and returns to the login page.",
        ],
    )

    add_heading(doc, "5. Data design and migrations")
    add_heading(doc, "5.1 Tables", 2)
    add_table(
        doc,
        ["Table", "Purpose", "Key relationships/indexes"],
        [
            ["waiver_submissions", "Parent waiver and signer record.", "Unique submission_reference; indexes on event, normalized contact data, event/email, and signed time."],
            ["waiver_family_members", "Zero-to-many family participants.", "Foreign key to submission with ON DELETE CASCADE; submission_id index."],
            ["email_outbox", "Durable customer and owner messages plus delivery state.", "Foreign key to submission with ON DELETE CASCADE; status/retry and creation indexes."],
        ],
        [1.45, 2.55, 2.5],
        font_size=8.6,
    )
    add_heading(doc, "5.2 Transaction boundary", 2)
    add_body(
        doc,
        "CreateSubmissionAsync inserts the submission first to obtain its numeric ID, then assigns "
        "that ID to every family member and outbox message. A single MySqlTransaction covers every "
        "insert. Any exception triggers rollback with CancellationToken.None so cancellation cannot "
        "leave a partial waiver.",
    )
    add_heading(doc, "5.3 Migration strategy", 2)
    add_bullets(
        doc,
        [
            "001_create_waiver_tables.sql creates the three permitted tables for a new database.",
            "002_add_media_release_agreed.sql upgrades an existing database with a NOT NULL boolean after agreed.",
            "DEFAULT TRUE initializes historical submissions as agreed and supports inserts that omit the new column.",
            "Normal application inserts always pass MediaReleaseAgreed explicitly; the database default does not override a posted opt-out.",
            "Apply migration 002 once before deploying application code that selects or inserts media_release_agreed.",
        ],
    )

    add_heading(doc, "6. Security and validation design")
    add_table(
        doc,
        ["Control", "Implementation"],
        [
            ["Server validation", "Both PageModel and WaiverSubmissionService validate; service validation remains authoritative."],
            ["SQL injection prevention", "All values use Dapper parameters. Admin sort columns map from an enum to fixed SQL fragments."],
            ["Cross-site request forgery", "Razor Pages form posts use the framework anti-forgery token."],
            ["Output encoding", "Razor encodes normal output; email bodies use HtmlEncoder for customer-supplied values."],
            ["Authentication", "Cookie authentication protects /Admin; cookie is HTTP-only, SameSite Strict, sliding, and eight hours."],
            ["Credential comparison", "Credentials are SHA-256 hashed before fixed-time byte comparison."],
            ["Open redirect prevention", "Login accepts ReturnUrl only when Url.IsLocalUrl returns true."],
            ["Secret handling", "Configuration supplies database, admin, and SMTP secrets; code avoids logging those values."],
            ["Request metadata", "IP and user agent are length-limited; user agent is truncated to 500 characters before mapping."],
        ],
        [1.55, 4.95],
        font_size=8.6,
    )

    add_heading(doc, "7. Email outbox design")
    add_body(
        doc,
        "Email is separated from the HTTP transaction. Submission success depends on queuing the "
        "messages, not on contacting SMTP. This prevents a temporary mail outage from losing or "
        "delaying the customer waiver.",
    )
    add_numbers(
        doc,
        [
            "EmailOutboxWorker starts only when Email:Enabled is true.",
            "The worker creates a dependency-injection scope and calls ProcessBatchAsync.",
            "GetPendingAsync selects Pending rows and Failed rows whose next attempt time has arrived.",
            "SmtpEmailSender validates settings and sends an HTML MailMessage.",
            "Success changes status to Sent and clears last_error.",
            "Failure stores a safe error type/summary and schedules exponential backoff of 1, 2, 4, 8, 16, 32, then at most 64 minutes.",
            "When the next attempt reaches MaximumAttempts, next_attempt_at_utc becomes null and the repository records Abandoned.",
        ],
    )
    add_callout(
        doc,
        "Operational behavior",
        "When no messages are eligible, the worker waits PollIntervalSeconds. Unexpected processor "
        "errors are logged and followed by the same delay. Application shutdown cancellation exits cleanly.",
    )

    add_heading(doc, "8. Administrator design")
    add_heading(doc, "8.1 Authentication", 2)
    add_body(
        doc,
        "Admin credentials are configured, not stored in MySQL. The login page is anonymous; the "
        "submissions page is authorized. Successful sign-in creates a name claim under the cookie scheme.",
    )
    add_heading(doc, "8.2 Query and pagination", 2)
    add_body(
        doc,
        "The repository first limits the dataset to the 200 newest submissions, then applies the "
        "selected sort and a 20-row LIMIT/OFFSET. This guarantees a maximum of ten pages and prevents "
        "older records from entering the admin browsing window merely because a different sort is selected.",
    )
    add_heading(doc, "8.3 Display projection", 2)
    add_body(
        doc,
        "AdminSubmissionListItem contains only fields needed by the table. A correlated GROUP_CONCAT "
        "subquery formats family names and optional relationships into one display string.",
    )

    add_heading(doc, "9. Configuration and deployment")
    add_table(
        doc,
        ["Configuration", "Purpose"],
        [
            ["ConnectionStrings:Default", "MySQL connection used by all repositories."],
            ["Waiver:BusinessOwnerEmail", "Recipient of every owner notification."],
            ["Waiver:Events", "Allowed event-code to event-name dictionary."],
            ["Email:Enabled", "Starts or disables the background delivery worker."],
            ["Email:FromAddress / FromName", "Sender identity shown on outgoing messages."],
            ["Email:PollIntervalSeconds", "Idle/error polling delay, clamped to 1-300 seconds."],
            ["Email:BatchSize", "Messages processed per batch, clamped to 1-100."],
            ["Email:MaximumAttempts", "Delivery attempts before abandonment, clamped to 1-20."],
            ["Email:Smtp:*", "SMTP host, port, TLS flag, username, and password."],
            ["Admin:Username / Password", "Credentials for the protected administrator area."],
        ],
        [2.25, 4.25],
        font_size=8.6,
    )
    add_heading(doc, "9.1 Deployment sequence", 2)
    add_numbers(
        doc,
        [
            "Provision MySQL and restrict network/database credentials appropriately.",
            "Apply 001 for a new database, or apply outstanding numbered migrations in order.",
            "Provide connection string, events, business-owner email, admin credentials, and SMTP settings through Railway variables or secure configuration.",
            "Build the Docker image and deploy the ASP.NET Core application.",
            "Verify the public event URL, one submission, both outbox rows, email delivery, admin login, and admin display.",
            "Monitor application logs for outbox failures without logging recipient PII or secrets.",
        ],
    )

    add_heading(doc, "10. Testing and maintenance")
    add_heading(doc, "10.1 Automated coverage", 2)
    add_bullets(
        doc,
        [
            "WaiverSubmissionService tests cover normalization, validation, duplicate acceptance, HTML encoding, media opt-out, emails, and owner configuration.",
            "Index PageModel tests cover event/agreement loading, form mapping, required agreement, and confirmation redirect.",
            "Repository SQL and migration tests verify explicit parameter use, required tables, keys, indexes, and media column.",
            "Email processor tests cover successful delivery, retry scheduling, cancellation, and abandonment.",
            "Credential and admin/page tests cover authentication and display behavior.",
            "MySQL integration tests verify commit, rollback, generated IDs, relationships, and UTC mapping when WAIVERAPP_TEST_MYSQL_CONNECTION is configured.",
        ],
    )
    add_heading(doc, "10.2 Routine maintenance checklist", 2)
    add_bullets(
        doc,
        [
            "Before code changes, review PROJECT.md and AGENTS.md.",
            "Keep agreement edits limited to Content/waiver-agreement.html and obtain business/legal approval.",
            "Add numbered migrations for schema changes; never assume CREATE TABLE IF NOT EXISTS alters an existing table.",
            "Update model, request, service, repository SQL, emails/admin projection, and tests together when adding a submission field.",
            "Run dotnet build and the full test suite; configure MySQL integration tests for release-critical database changes.",
            "Confirm email HTML in common clients when changing markup or inline styles.",
            "Do not log secrets, connection strings, passwords, or unnecessary customer data.",
        ],
    )
    add_heading(doc, "10.3 Known maintenance considerations", 2)
    add_callout(
        doc,
        "Agreement quality",
        "FileWaiverAgreementProvider considers any nonblank file approved unless it contains "
        "REPLACE_WITH_APPROVED_WAIVER. Legal approval and text-encoding quality remain operational responsibilities.",
        fill=AMBER,
    )
    add_bullets(
        doc,
        [
            "System.Net.Mail SMTP behavior should be revalidated if hosting or TLS requirements change.",
            "The simple configured admin credential model is suitable for the current small protected area; multiple users, password rotation, audit trails, or MFA would require a new authentication design.",
            "The outbox query does not claim/lock rows for multiple worker instances. Scale-out deployment would require row claiming or distributed coordination to prevent duplicate sends.",
            "Media consent defaults to selected and historical migration values default to true; changes to consent policy require coordinated UI, schema, wording, and legal/privacy review.",
        ],
    )

    add_heading(doc, "11. Function reference")
    add_table(
        doc,
        ["Component / function", "What it does"],
        [
            ["IndexModel.OnGetAsync", "Resolves the event and loads agreement content for the public form."],
            ["IndexModel.OnPostAsync", "Validates page state, maps input, submits, stores TempData, and redirects."],
            ["WaiverSubmissionService.FindEvent", "Normalizes a code and resolves it from configured allowed events."],
            ["WaiverSubmissionService.SubmitAsync", "Performs authoritative validation, constructs all records, and invokes transactional persistence."],
            ["ValidateFamilyMembers", "Enforces the ten-member limit and validates each nested entry."],
            ["CreateOutboxMessages", "Creates the customer confirmation and owner notification."],
            ["CreateBossNotificationBody", "Builds HTML-safe details including event highlight and media choice."],
            ["WaiverSubmissionRepository.CreateSubmissionAsync", "Commits submission, members, and emails in one MySQL transaction."],
            ["EmailOutboxRepository.GetPendingAsync", "Returns pending and due-for-retry messages in creation order."],
            ["EmailOutboxProcessor.ProcessBatchAsync", "Sends a batch and records sent, retry, or abandoned state."],
            ["EmailOutboxWorker.ExecuteAsync", "Runs the background polling loop and creates a scope per batch."],
            ["SmtpEmailSender.SendAsync", "Validates SMTP configuration and sends one HTML message."],
            ["AdminCredentialValidator.IsValid", "Checks both configured credentials using fixed-time comparisons."],
            ["LoginModel.OnPostAsync", "Authenticates the administrator and safely redirects."],
            ["SubmissionsModel.OnGetAsync", "Parses list controls, counts records, clamps pagination, and loads one page."],
            ["AdminSubmissionRepository.GetRecentAsync", "Queries a safe sort column and aggregates family display data."],
            ["FileWaiverAgreementProvider.GetAsync", "Reads the HTML agreement and reports whether submissions may proceed."],
            ["ConfirmationModel.OnGet", "Consumes one-time confirmation data and blocks direct/stale access."],
            ["MySqlConnectionFactory.OpenConnectionAsync", "Creates an open UTC-aware MySQL connection and cleans up on failure."],
        ],
        [2.6, 3.9],
        font_size=8.2,
    )
    doc.save(DESIGN_PATH)


if __name__ == "__main__":
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    build_specification()
    build_design()
    print(SPEC_PATH)
    print(DESIGN_PATH)
